from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date

# Initialize the document
doc = Document()

# Define a function to add a heading with custom formatting
def add_heading(text, level):
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    run.font.size = Pt(14)
    run.font.name = 'Arial'
    return heading

# Add title
title = doc.add_heading(level=1)
run = title.add_run("Ashwani Sharma")
run.font.size = Pt(24)
run.bold = True
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Add contact information
contact_info = doc.add_paragraph()
contact_info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
contact_info.add_run("Github: ").bold = True
contact_info.add_run("AshwaniShar | ")
contact_info.add_run("LinkedIn: ").bold = True
contact_info.add_run("ashwanishar | ")
contact_info.add_run("Email: ").bold = True
contact_info.add_run("ashwanisharma9793@gmail.com | ")
contact_info.add_run("Phone: ").bold = True
contact_info.add_run("+91-9076676489")

# Add summary section
add_heading("Summary", level=1)
summary = doc.add_paragraph(
    "Enthusiastic and motivated data science graduate with a solid foundation in statistical analysis, "
    "machine learning, and programming languages such as Python and R. Academic projects and coursework "
    "have provided hands-on experience in data cleaning, exploratory data analysis, and model building. "
    "Possesses a strong aptitude for problem-solving and a keen interest in leveraging data-driven insights "
    "to drive business growth. Excellent communicator with a collaborative mindset, eager to contribute to "
    "innovative projects and expand knowledge in the field of data science."
)

# Add work experience section
add_heading("Work Experience", level=1)
work_experience = doc.add_paragraph()
work_experience.add_run("Trainee\n").bold = True
work_experience.add_run("May 2024 - present\n").italic = True
work_experience.add_run(
    "Passionate and dedicated data science trainee with a strong foundation in statistics, programming, "
    "and machine learning concepts. Currently undergoing intensive training in data science methodologies, "
    "including data preprocessing, predictive modeling, and data visualization. Eager to apply theoretical knowledge "
    "to real-world projects and gain hands-on experience in analyzing complex datasets. A quick learner with excellent "
    "problem-solving skills and a proactive approach to continuous learning. Enthusiastic about contributing to "
    "cutting-edge projects and making meaningful contributions to the field of data science."
)

# Add projects section
add_heading("Projects", level=1)
project = doc.add_paragraph()
project.add_run("Web Scraping\n").bold = True
project.add_run("Link: https://github.com/AshwaniShar/web_scrapping\n").italic = True
project.add_run(
    "I embarked on a journey to explore the vast landscape of data available on the web. With a passion for extracting "
    "valuable insights from online sources, I dove into the world of web scraping techniques and tools. Armed with Python "
    "and popular libraries like BeautifulSoup and Scrapy, I embarked on various scraping projects, ranging from gathering "
    "product information for market analysis to extracting news articles for sentiment analysis."
)

# Add education section
add_heading("Education", level=1)
education = doc.add_paragraph()
education.add_run("M.Tech. at Dr. APJ Abdul Kalam Technical University\n").bold = True
education.add_run("2022 - present\n")
education.add_run("B.Tech. at Dr. APJ Abdul Kalam Technical University\n").bold = True
education.add_run("2018 - 2022\n")
education.add_run("Class 12th at Uttar Pradesh State Board of High School and Intermediate Education\n").bold = True
education.add_run("2018\n")
education.add_run("Class 10th at Uttar Pradesh State Board of High School and Intermediate Education\n").bold = True
education.add_run("2016\n")

# Add certificates section
add_heading("Certificates", level=1)
certificates = doc.add_paragraph()
certificates.add_run("Docker, Python - IBM\n").bold = True
certificates.add_run("JAVA, SQL, Python - Hackerrank\n").bold = True
certificates.add_run("Software Engineering - JP Morgan Chase and Co.\n").bold = True
certificates.add_run("Data Science - Top Mentor\n").bold = True

# Add skills section
add_heading("Skills", level=1)
skills = doc.add_paragraph()
skills.add_run("Skills: ").bold = True
skills.add_run(
    "Python, Machine Learning, Statistics, SQL, MongoDb, Power BI, Pandas, Numpy, Seaborn, Matplotlib, R, Computer Vision, "
    "Generative AI, Deep learning, Tableau\n"
)
skills.add_run("Some Other Skills: ").bold = True
skills.add_run("Microsoft Excel, Java, Android, Linux, Agile, Scraping\n")

# Add last updated date
doc.add_paragraph(f"Last updated: {date.today().strftime('%B %d, %Y')}").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Save the document
doc.save("Ashwani_Sharma_CV.docx")
